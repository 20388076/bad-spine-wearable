'''
Created on Tue May  6 14:39:30 2025

@author: AXILLIOS
'''
import os
import json
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import scipy as sc
from numpy import unique
import time
import sys
from sklearn import preprocessing
from sklearn.metrics import accuracy_score
from sklearn.model_selection import train_test_split
from scipy.stats import mode, randint, uniform
from micromlgen import port

from sklearn.metrics import (accuracy_score, precision_score, recall_score, f1_score, 
                           roc_auc_score, confusion_matrix, ConfusionMatrixDisplay, classification_report, 
                           roc_curve, auc, RocCurveDisplay)
import matplotlib.ticker as mticker

import io
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Pt, Inches

# import warnings
# warnings.filterwarnings("ignore")

def cls():
    print(chr(27) + '[2J') 
def pause():
    input('PRESS ENTER TO CONTINUE.')
#------------------------------------------------------------
def tic():
    t1=float(time.time());
    return t1
#------------------------------------------------------------
def toc(t1,s):
    t2=float(time.time());dt=t2-t1;
    s1='time taken '+s 
    print('%s %e' % (s1,dt) )     
#---------------------------------------------------------
def RETURN():
    sys.exit()

#--------------------------------------------------------------------

def my_nchoosek(nFeats, nCombs):
    from itertools import combinations
    x = np.arange(nFeats)
    L = [c for i in range(nCombs + 1) for c in combinations(x, i)]
    return L

#--------------------------------------------------------------------

def best_map(y_true, y_pred):
    # Match predicted cluster labels to true labels
    labels = np.unique(y_true)
    new_pred = np.zeros_like(y_pred)
    for label in labels:
        mask = y_pred == label
        new_pred[mask] = mode(y_true[mask])[0]
    return new_pred
# -------------------------------------------------------------------------

def split_classes(X, y):
    class1=X[y==0]
    class2=X[y==1]
    class3=X[y==2]
    return class1, class2, class3 

def RFE_wrapper(X, y, fNames, classif, desiredFeatures):
    
    (class1,class2,class3)=split_classes(X,y)
    # RFE with LogReg
    from sklearn.feature_selection import RFE
    from sklearn.linear_model import LogisticRegression

    nFeats = np.size(X, 1)
    
    # feature extraction

    if (classif == 0):
        #from sklearn.ensemble import GradientBoostingClassifier
        #model=GradientBoostingClassifier()
        #model = LogisticRegression(solver='saga')
        from sklearn.svm import SVC
        model = SVC(kernel='rbf', gamma=0.5, C=1.0)
        modelName ='Logistic Regression'
    elif (classif == 1):
        from sklearn.ensemble import RandomForestRegressor
        model = RandomForestRegressor()
        modelName ='Random Forest Regressor'
    elif (classif == 2):
        from sklearn.ensemble import ExtraTreesClassifier
        model = ExtraTreesClassifier(n_estimators=20)
        modelName ='Extra Trees Classifier'
    elif (classif == 3):
        from sklearn.tree import DecisionTreeRegressor
        model = DecisionTreeRegressor(max_leaf_nodes=20)
        modelName ='Decision Tree Regressor'

    rfe = RFE(model)#,n_features_to_select=15)# have to update the sklearn library
    fit = rfe.fit(X, y)
    print('Num Features: %d\n' % fit.n_features_)
    print('Selected Features: %s\n' % fit.support_)
    print('Feature Ranking: %s\n' % fit.ranking_)
    z = fit.ranking_
    indx=[]
    for i in range(nFeats):
        if (int(z[i]) == 1):
            indx.append(i)
    indx=np.asarray(indx,int)        
        
    print('selected-feature indices: ', end='')
    print(indx, end='')
    print('\n. Selected feature names:  ', end='')
    print(fNames[indx])
    
    #fNames = fNames[indx]
    if(np.size(X,1)>=desiredFeatures):
        X=X[:,0:desiredFeatures]
        fNames=fNames[0:desiredFeatures]
        
    return X, fNames, modelName

# -------------------------------------------------------------------------

def loadData(path, files, index):

    data_list = []
    y_list = []
    fNames = []

    for label, file in enumerate(files):
        df = pd.read_csv(path + file)

        # Extract and select feature names on first file
        if not fNames:
            all_features = df.columns[1:].tolist()  # Exclude time column

        all_features = df.columns[1:].tolist()  # Exclude time column
        fNames = all_features
        
        # Drop first row (label row) and time column, subset features
        df = df.iloc[1:, 1:][fNames]
        df = df.astype(np.float32)

        data_list.append(df)
        y_list.append(np.full(df.shape[0], label))

    X = np.concatenate(data_list, axis=0)
    y = np.concatenate(y_list, axis=0)
    
    if index == 0:
        Xn = X[:,0:6]
        tag ='ONLY RAW DATA'
        
    elif index == 1:
        Xn = X
        tag ='ALL FEATURES'

    elif index == 2:
        column_indices = [76, 32, 42, 46, 70, 28, 63, 27, 45, 62]
        Xn = X[:, column_indices]
        tag ='RELIFF FEATURES 10 Best'
        '''
        0:Index 76: ENERGY_gyro y
        1:Index 32: gyro_z_window_max
        2:Index 42: MAD_acceleration x
        3:Index 46: MAD_gyro y
        4:Index 70: FFT_gyro y
        5:Index 28: gyro_y_window_mean
        6:Index 63: IQR_gyro x
        7:ndex 27: gyro_x_window_min
        8:Index 45: MAD_gyro x
        9:Index 62: IQR_acceleration z
        '''
    return Xn, y, np.array(fNames), tag

    
# -------------------------------------------------------------------------
def init_document(docx_filename):
    """Create a new document if it doesn't exist, otherwise load the existing one."""
    if os.path.exists(docx_filename):
        return Document(docx_filename)
    else:
        return Document()
        

#====================================================================
#========================== MAIN PROGRAM ============================
#====================================================================

cls() 
docx_filename = 'classification_results.docx'
doc = init_document(docx_filename)

files = ['x_axis_with _random_movements_feat.csv',
         'x_1deg_per_min_feat.csv', 
         'x_2deg_per_min_feat.csv',
         'x_anomaly_detection_3dpersec_feat.csv',
         'y_axis_with _random_movements_feat.csv',
         'y_1deg_per_min_feat.csv', 
         'y_2deg_per_min_feat.csv',
         'y_anomaly_detection_3dpersec_feat.csv',
         'z_axis_with _random_movements_feat.csv',
         'z_1deg_per_min_feat.csv', 
         'z_2deg_per_min_feat.csv',
         'z_anomaly_detection_3dpersec_feat.csv',
    ]
path = './FEATS/'

for choice in range(0,2):
    for index in range(0,3):
        X, y, fNames, Data_tag = loadData(path, files, index)
        
        #============================== FEATURE SELECTION =============================
        '''
        desiredFeatures=10
        
        classif = 3  # 0: for logistic regressor 1:for RF regressor 2:ExtraTreesClassifier
        # 3.DecisionTreeRegressor
        (Xn, fNames, modelName) = RFE_wrapper(X, y, fNames, classif, desiredFeatures)
        print('\nRFE_wrapper for %d best features with %s model\n'%(desiredFeatures, modelName))
        print(fNames)
        '''
        #============================== TRAIN/TEST SPLIT ==============================
    
        def TrainTestSplit(X, y, train_size, test_size):
    
            X_train_parts = []
            X_test_parts = []
            y_train_parts = []
            y_test_parts = []
    
            for class_label in np.unique(y):
                class_data = X[y == class_label]
                X_train_class, X_test_class = train_test_split(
                    class_data, train_size=train_size, test_size=test_size, shuffle=False
                )
                X_train_parts.append(X_train_class)
                X_test_parts.append(X_test_class)
                y_train_parts.append(np.full(X_train_class.shape[0], class_label))
                y_test_parts.append(np.full(X_test_class.shape[0], class_label))
    
            X_train = np.concatenate(X_train_parts, axis=0)
            X_test = np.concatenate(X_test_parts, axis=0)
            y_train = np.concatenate(y_train_parts, axis=0)
            y_test = np.concatenate(y_test_parts, axis=0)
    
            return X_train, X_test, y_train, y_test
        # -------------------------------------------------------------------------
        
        X_train, X_test, y_train, y_test = TrainTestSplit(X, y, train_size=0.7, test_size=0.3)
        
        #================================ CLASSIFICATION ==============================
        
        import time
        search = 2
        
        if search == 0:
            
            def classifiers(choice):
                #if choice == 0: # Nearest Centroid
                    #from sklearn.neighbors import KNeighborsClassifier
                    #return KNeighborsClassifier(), 'KNN'
                    
                if choice == 0: # Decision Tree 
                    from sklearn.tree import DecisionTreeClassifier  # Import Decision Tree Classifier
                    return DecisionTreeClassifier(), 'DecisionTree'
                    
                elif choice == 1: # Random Forest 
                    from sklearn.ensemble import RandomForestClassifier
                    return RandomForestClassifier(n_estimators=10), 'RandomForest'     
                    
                #elif choice == 2: # eXtreme Gradient Boosting
                    #from xgboost import XGBClassifier
                    #return XGBClassifier(n_estimators=10, learning_rate=0.1, max_depth=7), 'XGBoost'
            
            start = time.time()
            classifier, classifier_name = classifiers(choice)
            classifier = classifier.fit(X_train, y_train)
            end = time.time()
            classification_time = end-start
            y_pred = classifier.predict(X_test)
            
        elif search == 1: # Hyperparameters Tuning
        
            def classifiers(choice):
                #if choice == 0: # Nearest Centroid
                    #from sklearn.neighbors import KNeighborsClassifier
                    #return KNeighborsClassifier(), 'KNN'
                    
                if choice == 0: # Decision Tree 
                    from sklearn.tree import DecisionTreeClassifier  # Import Decision Tree Classifier
                    return DecisionTreeClassifier(), 'DecisionTree'
                    
                elif choice == 1: # Random Forest 
                    from sklearn.ensemble import RandomForestClassifier
                    return RandomForestClassifier(), 'RandomForest'     
                    
                #elif choice == 3: # eXtreme Gradient Boosting
                    #from xgboost import XGBClassifier
                    #return XGBClassifier(use_label_encoder=False, eval_metric='mlogloss'), 'XGBoost'
            
            start = time.time()
            classifier, classifier_name = classifiers(choice)
            classifier = classifier.fit(X_train, y_train)
            end = time.time()
            classification_time = end-start
            y_pred = classifier.predict(X_test)
            
            # Define parameter grids
            param_grids = {
                #'KNN': {'n_neighbors': [3, 5, 7, 9],'weights': ['uniform', 'distance'],'metric': ['euclidean', 'manhattan']},
                'DecisionTree': {
                    'max_depth': [None, 5, 10, 20],
                    'min_samples_split': [2, 5, 10]
                },
                'RandomForest': {
                    'n_estimators': [50, 100, 200],
                    'max_depth': [None, 10, 20],
                    'min_samples_split': [2, 5]
                },
                #'XGBoost': {'n_estimators': [50, 100, 200, 300],'max_depth': [3, 5, 7, 10],'learning_rate': [0.01, 0.1, 0.3, 1]}
            }
            
            # Parameter distributions for RandomizedSearchCV
            param_dists = {
                #KNN': {'n_neighbors': randint(3, 15),'weights': ['uniform', 'distance'],'metric': ['euclidean', 'manhattan']},
                'DecisionTree': {
                    'max_depth': randint(3, 20),
                    'min_samples_split': randint(2, 10)
                },
                'RandomForest': {
                    'n_estimators': randint(50, 200),
                    'max_depth': randint(3, 20),
                    'min_samples_split': randint(2, 10)
                }
                #XGBoost': {'n_estimators': randint(50, 300),'max_depth': randint(3, 10),'learning_rate': uniform(0.01, 1)}
            }
            
            from sklearn.model_selection import GridSearchCV, RandomizedSearchCV
            # GridSearchCV or RandomizedSearchCV
            search_type = 1  # 0 for grid or 1 for random
            # Store results
            summary = []
            if search_type == 0:
                search = GridSearchCV(classifier, param_grids[classifier_name], cv=5, scoring='accuracy')
            else:
                search = RandomizedSearchCV(classifier, param_dists[classifier_name], n_iter=20, cv=5, scoring='accuracy', random_state=42)
            # Fit and get results
            search.fit(X_train, y_train)
            best_params = search.best_params_
            best_cv_score = search.best_score_
            
            # Save parameters
            filename = f"best_params_{classifier_name}_{Data_tag.replace(' ', '_')}.json"
            with open(filename, "w") as f:
                json.dump(best_params, f)
        
            # Test set accuracy
            best_model = search.best_estimator_
            y_pred = best_model.predict(X_test)
            test_score = accuracy_score(y_test, y_pred)
        
            # Store for summary
            summary.append({
                'Classifier': classifier_name,
                'CV Accuracy': round(best_cv_score, 4),
                'Test Accuracy': round(test_score, 4),
                'Best Params': best_params,
                'Saved As': filename
            })
            # Summary
            print("\n Summary:")
            for s in summary:
                print(f"\n {s['Classifier']}")
                print(f"   CV Accuracy  : {s['CV Accuracy']}")
                print(f"   Test Accuracy: {s['Test Accuracy']}")
                print(f"   Best Params  : {s['Best Params']}")
                print(f"   Saved To     : {s['Saved As']}")
        elif search == 2:
            
            if choice == 0:
                classifier_name = 'DecisionTree'
            elif choice == 1:  
                classifier_name = 'RandomForest'
            
            def load_classifier(choice, params):

                if choice == 0:
                    from sklearn.tree import DecisionTreeClassifier
                    return DecisionTreeClassifier(**params)
                elif choice == 1:
                    from sklearn.ensemble import RandomForestClassifier
                    return RandomForestClassifier(**params)
            
            json_file = f"best_params_{classifier_name}_{Data_tag.replace(' ', '_')}.json"
            try:
                with open(json_file, "r") as f:
                    params = json.load(f)
                classifier = load_classifier(choice, params)
                start = time.time()
                classifier.fit(X_train, y_train)
                end = time.time()
                classification_time = end-start
                y_pred = classifier.predict(X_test)
            except FileNotFoundError:
                print(f"  Could not find '{json_file}'")                    
                    
            # Export model to c code
            model_code = port(classifier)
            
            # Windows path (use raw string or double backslashes)
            save_dir = r"C:\Users\user\OneDrive\Έγγραφα\PlatformIO\Projects\bad-spine-wearable-1\src"
            
            # Ensure path exists (optional safety check)
            os.makedirs(save_dir, exist_ok=True)
            
            # Full path to save header file
            save_path = os.path.join(save_dir, f"{classifier_name}_{Data_tag.replace(' ', '_')}.h")
            
            # Write the file
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(model_code)
                
        #========================== PLOTING AND WORD SAVING ==========================

        accuracy = accuracy_score(y_test, y_pred) * 100
        acc = f" Accuracy: {accuracy:.2f} % "
        #accuracy = print_confusion_matrix(y_test, y_pred)
        print('=' * 23 + ' CLASSIFICATION RESULTS ' + '=' * 23)
        print('CLASSIFIER: ', classifier_name)
        print(acc)
        print('Feature used :', Data_tag)
        print('Classification Time: ', np.asarray(classification_time).round(3),'seconds')
        
        
        # Helper to simulate console output and capture it
        def capture_output_and_plot(classifier_name, accuracy, Data_tag, classification_time, classifier, X_test, y_test):
            buf = io.StringIO()
            with redirect_stdout(buf):
                print('=' * 23 + ' CLASSIFICATION RESULTS ' + '=' * 23)
                print('CLASSIFIER: ', classifier_name)
                print(acc)
                print('Feature used :', Data_tag)
                print('Classification Time: ', np.asarray(classification_time).round(3),'seconds')
                # Plot confusion matrix image as an example 
                # Original class names from filenames
                original_class_names = [filename.replace('_feat.csv', '') for filename in files]
                
                # Generate numeric labels for display
                class_names = [f'Class {i}' for i in range(len(original_class_names))]
                
                # Create a mapping legend
                legend_text = '\n'.join([f'{class_names[i]}: {original_class_names[i]}' for i in range(len(class_names))])
                plt.figure(figsize=(13.5,10))
                # Plot confusion matrix
                disp = ConfusionMatrixDisplay.from_estimator(
                    classifier,
                    X_test,
                    y_test,
                    display_labels=class_names,
                    cmap=plt.cm.Blues,
                    xticks_rotation=90, 
                    values_format='d'
                )
                
                # Remove scientific notation    
                plt.gca().set_xticklabels(class_names, rotation=90)
                plt.gca().set_yticklabels(class_names)
                
                # Add the legend as a textbox
                plt.gcf().text(1.02, 
                               0.5, 
                               legend_text, 
                               fontsize=8, 
                               va='center', 
                               bbox=dict(facecolor='white', edgecolor='black')
                               )
               
                # Title
                plt.title(f'{classifier_name} Confusion Matrix ' + Data_tag 
                          + acc , fontsize=16)
                
                # Improve readability
                plt.tick_params(axis='x', labelsize=10)
                plt.tick_params(axis='y', labelsize=10)
                
                # Optional: Bold larger numbers or set font size
                for text in disp.ax_.texts:
                    text.set_fontsize(6)  # Increase for better visibility (try 10–12 if need 
                # Save and show
                plt.savefig('temp_image.png', dpi=600, bbox_inches='tight')
                plt.show()
                plt.close()
        
            return buf.getvalue(), 'temp_image.png'
        
        # Capture simulated console output and image
        word = 1
        
        if word == 0:
            classification_text, image_path = capture_output_and_plot(classifier_name, accuracy, Data_tag, classification_time, classifier, X_test, y_test)
        else:
            classification_text, image_path = capture_output_and_plot(classifier_name, accuracy, Data_tag, classification_time, classifier, X_test, y_test)
            def append_results_to_doc(doc, classification_text, image_path, font_name='Times New Roman', font_size=11):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(classification_text)
            
                # Set font properties
                font = run.font
                font.name = font_name
                font.size = Pt(font_size)
            
                # Add image
                max_width_in_inches = 6.0  # You can change this to fit your layout
                doc.add_picture(image_path, width=Inches(max_width_in_inches))
                doc.add_paragraph('\n' + '='*70 + '\n')  # Add separation between results
                
      
            append_results_to_doc(
                doc,
                classification_text,
                image_path,
                font_name='Times New Roman',
                font_size=11
                )
 
doc.save(docx_filename)
    
  
                
'''
with open(f"{classifier_name}.h", "w") as f:
    f.write(port(model)) 
# NOT WORKING PROPERLY
def roc_curve_display(classifier, y_train, Xn_test, y_test, class_names):
from itertools import cycle
from sklearn.preprocessing import LabelBinarizer
# Binarize the test labels
label_binarizer = LabelBinarizer().fit(y_train)
y_onehot_test = label_binarizer.transform(y_test)

# Checking if there is binary classification
if y_onehot_test.shape[1] == 1:
    y_onehot_test = np.hstack((1 - y_onehot_test, y_onehot_test))

# Get predicted probabilities
y_pred_proba = classifier.predict_proba(Xn_test)

# ROC plotting
fig, ax = plt.subplots(figsize=(6, 6))
colors = cycle(['aqua', 'darkorange', 'cornflowerblue'])
n_classes = y_onehot_test.shape[1]

for class_id, color in zip(range(n_classes), colors):
    RocCurveDisplay.from_predictions(
        y_true=y_onehot_test[:, class_id],
        y_pred=y_pred_proba[:, class_id],
        name=f'ROC curve for {class_names[class_id]}',
        color=color,
        ax=ax,
        #plot_chance_level=(class_id == (n_classes - 1)),  # Plot chance level only once
    )

ax.set(title='ROC Curve', xlabel='False Positive Rate', ylabel='True Positive Rate')
plt.grid(True)
plt.show()
return

roc_curve_display(classifier, y_train, Xn_test, y_test, class_names)
'''


'''
def roc_curve_display(y_train,X_test,y_test):
    from sklearn.preprocessing import LabelBinarizer
    label_binarizer = LabelBinarizer().fit(y_train)
    y_onehot_test = label_binarizer.transform(y_test)
    y_pred_proba = classifier.predict_proba(X_test)
    from itertools import cycle
    fig, ax = plt.subplots(figsize=(6, 6))
    colors = cycle(['aqua', 'darkorange', 'cornflowerblue'])
    n_classes = 3
    for class_id, color in zip(range(n_classes), colors):
        RocCurveDisplay.from_predictions(
            y_onehot_test[:, class_id],
            y_pred_proba[:, class_id],
            name = f'ROC curve for {class_names[class_id]}',
            color = color,
            ax = ax,
            plot_chance_level = (class_id == 2),
            despine = True,
        )
    _ = ax.set(title = 'ROC curve')
    return
disp = roc_curve_display(y_train,X_test,y_test)
'''

'''
import csv

data1 = X.tolist()
data2 = y.tolist()
file_path1 = 'data_X.csv'
file_path2 = 'data_y.csv'

# Save data1 (X) — assume it's a 2D array
with open(file_path1, 'w', newline='') as file:
    writer = csv.writer(file)
    writer.writerows(data1)

# Save data2 (y) — assume it's a 1D array
with open(file_path2, 'w', newline='') as file:
    writer = csv.writer(file)
    for val in data2:
        writer.writerow([val])  # Wrap in list to write one value per row 

import csv
data1 = Xn.tolist()
data2 = y.tolist()
file_path1 = 'data_Xn.csv'
file_path2 = 'data_yn.csv'

# Save data1 (X) — assume it's a 2D array
with open(file_path1, 'w', newline='') as file:
    writer = csv.writer(file)
    writer.writerows(data1)

# Save data2 (y) — assume it's a 1D array
with open(file_path2, 'w', newline='') as file:
    writer = csv.writer(file)
    for val in data2:
        writer.writerow([val])  # Wrap in list to write one value per row 
        
from skrebate import ReliefF

# ReliefF Feature Selection
relieff = ReliefF(n_neighbors=3, n_features_to_select=5)
relieff.fit(X_train, y_train)

# Get top selected features
top_k = 3
selected_indices = relieff.top_features_[:top_k]
selected_feature_names = fNames[selected_indices]

print('Top', top_k, 'features selected by ReliefF:')
print(selected_feature_names)


fs = ReliefF(n_neighbors=5, n_features_to_keep=7)
X_train = fs.fit_transform(Xn, y)
print(X_train)

from sklearn.feature_selection import SelectKBest, f_classif

selector = SelectKBest(score_func=f_classif, k=10)
X_new = selector.fit_transform(X, y)

selected_indices = selector.get_support(indices=True)
selected_feature_names = fNames[selected_indices]

print('Top 10 features by ANOVA F-test:')
print(selected_feature_names)
'''
