'''
Created on Tue May  6 14:39:30 2025

@author: AXILLIOS
'''
# -------------------------------  Working Directory -----------------------------
# Set the working directory to the script's location if running in Visual Studio Code
import os
# Change working directory for this script
os.chdir(r'C:\Users\user\OneDrive\Έγγραφα\PlatformIO\Projects\bad-spine-wearable-1\src\data_process_and_classification') # modify this path to your working directory

# ============================= Utility Functions =============================

# ----------------------------- Import Libraries ------------------------------
import time
import sys
# ----------------------------- Kernel clean ----------------------------------
def cls():
    print(chr(27) + '[2J') 
# ----------------------------- Kernel pause ----------------------------------
def pause():
    input('PRESS ENTER TO CONTINUE.')
# ----------------------------- Process time count ----------------------------
def tic():
    return float(time.time())
# ----------------------------- Process time return ---------------------------
def toc(t1, s):
    t2 = float(time.time())
    print(f'{s} time taken: {t2 - t1:.6e} seconds')
# ----------------------------- Kernel break ----------------------------------
def RETURN():
    sys.exit()
# =============================================================================

# ----------------------------- Kernel clean call -----------------------------
cls()

import os
import json
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from numpy import unique
from sklearn import preprocessing
from sklearn.model_selection import train_test_split, TimeSeriesSplit, cross_validate
from scipy.stats import mode, randint, uniform
from micromlgen import port

from sklearn.metrics import (accuracy_score, precision_score, recall_score, f1_score, 
                             roc_auc_score, confusion_matrix, ConfusionMatrixDisplay, classification_report, 
                             roc_curve, auc, RocCurveDisplay, make_scorer, root_mean_squared_error)

import io
from contextlib import redirect_stdout
from docx import Document
from docx.shared import Pt, Inches

# import warnings
# warnings.filterwarnings('ignore')

def init_document(docx_filename):
    '''Create a new document if it doesn't exist, otherwise load the existing one.'''
    if os.path.exists(docx_filename):
        return Document(docx_filename)
    else:
        return Document()

#====================================================================
#========================== MAIN PROGRAM ============================
#====================================================================

def lower_bound(cv_results):
    r"""
    Calculate the lower bound within 1 standard deviation
    of the best `mean_test_scores`.

    Parameters
    ----------
    cv_results : dict of numpy(masked) ndarrays
        See attribute cv_results_ of `GridSearchCV`

    Returns
    -------
    float
        Lower bound within 1 standard deviation of the
        best `mean_test_score`.
    """
    best_score_idx = np.argmax(cv_results["mean_test_score"])

    return (
        cv_results["mean_test_score"][best_score_idx]
        - cv_results["std_test_score"][best_score_idx]
    )

def best_low_complexity(cv_results):
    r"""
    Balance model complexity with cross-validated score.

    Parameters
    ----------
    cv_results : dict of numpy(masked) ndarrays
        See attribute cv_results_ of `GridSearchCV`.

    Return
    ------
    int
        Index of a model that has the fewest PCA components
        while has its test score within 1 standard deviation of the best
        `mean_test_score`.
    """
    threshold = lower_bound(cv_results)
    candidate_idx = np.flatnonzero(cv_results["mean_test_score"] >= threshold)
    best_idx = candidate_idx[
        cv_results["param_reduce_dim__n_components"][candidate_idx].argmin()
    ]
    return best_idx

def loadData(paths, index, input_file):
    r'''
    Load dataset X, y, feature names, and tag based on index from input_file.
    
    Parameters
    ----------
    paths : list
        [stage4_path, stage6_path]
    index : int
        Row index of input_file
    
    Returns
    -------
    X : np.ndarray
    y : np.ndarray
    fNames : list
        Feature names
    Data_tag : str
        Tag describing the dataset
    '''

    row = input_file.iloc[index]
    X_file, y_file, Data_tag = row['X_file'], row['y_file'], row['Data_tag']

    # First 3 datasets are from Stage 4, rest from Stage 6
    base_path = paths[0] if index < 3 else paths[1]

    X_path = os.path.join(base_path, X_file)
    y_path = os.path.join(base_path, y_file)

    # ---- Load X ----

    X_data = pd.read_csv(X_path, dtype=np.float32)

    fNames = X_data.columns.tolist()
    X = X_data.to_numpy(dtype=np.float32)

    # ---- Load y ----
    y_data = pd.read_csv(y_path, header=None)

    # If first row is text (e.g., 'label'), drop it
    if isinstance(y_data.iloc[0, 0], str):
        y_data = y_data.iloc[1:]

    y = y_data.squeeze('columns').astype(np.int32).to_numpy()

    return X, y, fNames, Data_tag

cls() 
docx_filename = 'classification_results.docx'
doc = init_document(docx_filename)

files = ['movement_0_feat.csv',
         'x_axis_with_random_movements_feat.csv',
         'x_1step_per_min_feat.csv', 
         'x_2step_per_min_feat.csv',
         'x_anomaly_detection_3_step_per_min_feat.csv',
         'y_axis_with_random_movements_feat.csv',
         'y_1step_per_min_feat.csv', 
         'y_2step_per_min_feat.csv',
         'y_anomaly_detection_3_step_per_min_feat.csv',
         'z_axis_with_random_movements_feat.csv',
         'z_1step_per_min_feat.csv', 
         'z_2step_per_min_feat.csv',
         'z_anomaly_detection_3_step_per_min_feat.csv',
    ]
paths = ['./4_FEATS_COMBINED/','./5_FEATS_SELECTION/']

input_file_train = pd.DataFrame([
    ["all_raw_train.csv", "y_all_raw_train.csv", "RAW_DATA"],
    ["all_norm_train.csv", "y_all_norm_train.csv", "G_RAW_DATA"],
    ["X_train.csv", "y_train.csv", "ALL FEATURES"],
    ["Matlab_X_train_top10.csv", "y_train.csv", "MATLAB RELIFF FEATURES 10 Best"],
    ["Matlab_Top10_w_comp_efficient.csv", "y_train.csv", "MATLAB RELIFF Weight/Computation time 10 Best scored"]
        ], columns=['X_file', 'y_file', 'Data_tag'])
    
input_file_test = pd.DataFrame([
    ["all_raw_test.csv", "y_all_test.csv", "RAW_DATA"],
    ["all_norm_test.csv", "y_all_norm_test.csv", "G_RAW_DATA"],
    ["X_test.csv", "y_test.csv", "ALL FEATURES"],
    ["Matlab_X_test_top10.csv", "y_test.csv", "MATLAB RELIFF FEATURES 10 Best"],
    ["Matlab_Top10_w_comp_efficient.csv", "y_train.csv", "MATLAB RELIFF Weight/Computation time 10 Best scored"]
        ], columns=['X_file', 'y_file', 'Data_tag'])

import itertools

for cl, index in itertools.product( range(2), range(5)):
        
        X, y, fNames, Data_tag_train = loadData(paths, index, input_file_train)
        
        X_test, y_test, fNames, Data_tag_test = loadData(paths, index, input_file_test)
        
        #============================== TRAIN/TEST SPLIT ==============================
        
        def TrainTestSplit(X, y, train_size=0.75, test_size=0.25):
            '''
            Deterministic split per class: keeps order, no shuffle.
            Train = first part of rows
            Test  = last part of rows
            '''
            # Ensure y is a Series with same index
            if isinstance(y, np.ndarray):
                y = pd.Series(y, index=X.index, name='label')
            elif isinstance(y, pd.DataFrame):
                # flatten single-column DataFrame to Series
                y = y.iloc[:, 0]

            X_train_parts, X_test_parts = [], []
            y_train_parts, y_test_parts = [], []

            for class_label in y.unique():
                class_mask = (y == class_label)
                class_data = X.loc[class_mask]
                class_labels = y.loc[class_mask]

                n_total = len(class_data)
                n_train = int(n_total * train_size)

                # Deterministic split (no shuffle)
                X_train_class = class_data.iloc[:n_train]
                X_test_class  = class_data.iloc[n_train:]
                y_train_class = class_labels.iloc[:n_train]
                y_test_class  = class_labels.iloc[n_train:]

                # Collect
                X_train_parts.append(X_train_class)
                X_test_parts.append(X_test_class)
                y_train_parts.append(y_train_class)
                y_test_parts.append(y_test_class)

            # Concatenate back
            X_train = pd.concat(X_train_parts, axis=0, ignore_index=True)
            X_test  = pd.concat(X_test_parts, axis=0, ignore_index=True)
            y_train = pd.concat(y_train_parts, axis=0, ignore_index=True)
            y_test  = pd.concat(y_test_parts, axis=0, ignore_index=True)

            # Safeguard: total lengths must match input
            assert len(X_train) + len(X_test) == len(X), \
                f'Split mismatch: {len(X)} rows in, {len(X_train)+len(X_test)} out'
            assert len(y_train) + len(y_test) == len(y), \
                f'Label mismatch: {len(y)} rows in, {len(y_train)+len(y_test)} out'

            return X_train, X_test, y_train, y_test

        # -------------------------------------------------------------------------
        
        X_train, X_validation, y_train, y_validation = TrainTestSplit(X, y, train_size=0.7, validation_size=0.3)
        
        #================================ CLASSIFICATION ==============================
        
        search = 0
          
        if search == 0: # Hyperparameters Tuning
        
            def classifiers(cl):
                    
                if cl == 0: # Decision Tree 
                    from sklearn.tree import DecisionTreeClassifier  # Import Decision Tree Classifier
                    return DecisionTreeClassifier(), 'DecisionTree'
                    
                elif cl == 1: # Random Forest 
                    from sklearn.ensemble import RandomForestClassifier
                    return RandomForestClassifier(), 'RandomForest'     
                    
            
            start = time.time()
            classifier, classifier_name = classifiers(cl)
            classifier = classifier.fit(X_train, y_train)
            end = time.time()
            classification_time = end-start
            y_pred = classifier.predict(X_validation)
            
            # Define parameter grids
            param_grids = {
                #'KNN': {'n_neighbors': [3, 5, 7, 9],'weights': ['uniform', 'distance'],'metric': ['euclidean', 'manhattan']},
                'DecisionTree': {
                    'max_depth': [None, 5, 10, 20],
                    'min_samples_split': [2, 5, 10]
                },
                'RandomForest': {
                    'n_estimators': [5, 50, 100],
                    'max_depth': [None, 10, 20],
                    'min_samples_split': [2, 5]
                },
                #'XGBoost': {'n_estimators': [50, 100, 200, 300],'max_depth': [3, 5, 7, 10],'learning_rate': [0.01, 0.1, 0.3, 1]}
            }
            
            # Parameter distributions for RandomizedSearchCV
            param_dists = {
                #KNN': {'n_neighbors': randint(3, 15),'weights': ['uniform', 'distance'],'metric': ['euclidean', 'manhattan']},
                'DecisionTree': {
                    'max_depth': randint(1, 20),
                    'min_samples_split': randint(2, 10)
                },
                'RandomForest': {
                    'n_estimators': randint(1, 100),
                    'max_depth': randint(1, 20),
                    'min_samples_split': randint(2, 10)
                }
                #XGBoost': {'n_estimators': randint(50, 300),'max_depth': randint(3, 10),'learning_rate': uniform(0.01, 1)}
            }
            
            ts_cv = TimeSeriesSplit(n_splits=5) 
            from sklearn.model_selection import GridSearchCV, RandomizedSearchCV
            # GridSearchCV or RandomizedSearchCV
            search_type = 1  # 0 for grid or 1 for random
            # Store results
            summary = []
            if search_type == 0:
                search = GridSearchCV(classifier, param_grids[classifier_name], cv=ts_cv, scoring='accuracy')
            else:
                search = RandomizedSearchCV(classifier, param_dists[classifier_name], 
                                            n_iter=50, cv=ts_cv, n_jobs=-1, scoring='accuracy', random_state=42)
            # Fit and get results
            search.fit(X_train, y_train)
            best_params = search.best_params_
            best_cv_score = search.best_score_
            
            # Save parameters
            filename = f'best_params_{classifier_name}.json'
            with open(filename, 'w') as f:
                json.dump(best_params, f)
        
            # Test set accuracy
            best_model = search.best_estimator_
            y_pred = best_model.predict(X_validation)
            test_score = accuracy_score(y_validation, y_pred)
        
            # Store for summary
            summary.append({
                'Classifier': classifier_name,
                'CV Accuracy': round(best_cv_score, 4),
                'Test Accuracy': round(test_score, 4),
                'Best Params': best_params,
                'Saved As': filename
            })
            # Summary
            print('\n Summary:')
            for s in summary:
                print(f'\n {s['Classifier']}')
                print(f'   CV Accuracy  : {s['CV Accuracy']}')
                print(f'   Test Accuracy: {s['Test Accuracy']}')
                print(f'   Best Params  : {s['Best Params']}')
                print(f'   Saved To     : {s['Saved As']}')
                
            word = 0   
            
        elif search == 1:
            
            if cl == 0:
                classifier_name = 'DecisionTree'
            elif cl == 1:  
                classifier_name = 'RandomForest'
            
            def load_classifier(cl, params):

                if cl == 0:
                    from sklearn.tree import DecisionTreeClassifier
                    return DecisionTreeClassifier(**params)
                elif cl == 1:
                    from sklearn.ensemble import RandomForestClassifier
                    return RandomForestClassifier(**params)
            
            json_file = f'best_params_{classifier_name}.json'
            try:
                with open(json_file, 'r') as f:
                    params = json.load(f)
                classifier = load_classifier(cl, params)
                start = time.time()
                classifier.fit(X_train, y_train)
                end = time.time()
                classification_time = end-start
                y_pred = classifier.predict(X_test)
            except FileNotFoundError:
                print(f"  Could not find '{json_file}'")                    
                    
            # Export model to c code
            model_code = port(classifier)
            
            # Windows path (use raw string or double backslashes)
            save_dir = '../' 
            
            # Ensure path exists (optional safety check)
            os.makedirs(save_dir, exist_ok=True)
            
            # Full path to save header file
            save_path = os.path.join(save_dir, f'{classifier_name}.h')
            
            # Write the file
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(model_code)
                
            word = 1
                
        #========================== PLOTING AND WORD SAVING ==========================

        accuracy = accuracy_score(y_test, y_pred) * 100
        acc = f'Accuracy: {accuracy:.2f} % '
        #accuracy = print_confusion_matrix(y_test, y_pred)
        print('=' * 23 + ' CLASSIFICATION RESULTS ' + '=' * 23)
        print('CLASSIFIER: ', classifier_name)
        print(acc)
        print('Feature used :', Data_tag)
        print('Classification Time: ', np.asarray(classification_time).round(3),'seconds')
        print(f'Test set X__test shape: {X_test.shape}')
        
        # Helper to simulate console output and capture it
        def capture_output_and_plot(classifier_name, accuracy, Data_tag, 
                                    classification_time, classifier, X_test, y_test):
            buf = io.StringIO()
            with redirect_stdout(buf):
                print('=' * 23 + ' CLASSIFICATION RESULTS ' + '=' * 23)
                print('CLASSIFIER: ', classifier_name)
                print(acc)
                print('Feature used :', Data_tag)
                print('Classification Time: ', np.asarray(classification_time).round(3),'seconds')
                # Plot confusion matrix image as an example 
                # Original class names from filenames
                original_class_names = [filename.replace('_feat.csv', '') 
                                        for filename in files]
                
                # Generate numeric labels for display
                class_names = [f'Class {i}' for i in range(len(original_class_names))]
                
                # Create a mapping legend
                legend_text = '\n'.join([f'{class_names[i]}: {original_class_names[i]}' 
                                         for i in range(len(class_names))])
                
                labels = np.unique(y_test)  # integers that appear in y_test
                fig, ax = plt.subplots(figsize=(10, 8))  # wider & taller
                disp = ConfusionMatrixDisplay.from_estimator(
                    classifier,
                    X_test,
                    y_test,
                    labels=labels,                                      
                    display_labels=[f'Class {i}' for i in labels],       
                    normalize='true',
                    cmap=plt.cm.Blues,                                    # row-wise %
                    xticks_rotation=90,
                    values_format='.1%',  
                    ax=ax            # show as percents
                    )
                
                # Remove scientific notation    
                plt.gca().set_xticklabels(class_names, rotation=90)
                plt.gca().set_yticklabels(class_names)
                
                # Add the legend as a textbox
                plt.gcf().text(1.02, 
                               0.5, 
                               legend_text, 
                               fontsize=12, 
                               va='center', 
                               bbox=dict(facecolor='white', edgecolor='black')
                               )
               
                # Title
                plt.title(f'{classifier_name} Confusion Matrix \n' + Data_tag 
                          + acc , fontsize=16)
                
                # Improve readability
                plt.tick_params(axis='x', labelsize=10)
                plt.tick_params(axis='y', labelsize=10)
                
                # Optional: Bold larger numbers or set font size
                for text in disp.ax_.texts:
                    text.set_fontsize(6)  # Increase for better visibility (try 10–12 if need 
                # Save and show
                plt.savefig('temp_image.png', dpi=600, bbox_inches='tight')
                plt.tight_layout()
                plt.show()
                plt.close()
        
            return buf.getvalue(), 'temp_image.png'
        
        
        # Capture simulated console output and image
        
        
        if word == 0:
            classification_text, image_path = capture_output_and_plot(classifier_name,
                                                                      accuracy, Data_tag, 
                                                                      classification_time, 
                                                                      classifier, 
                                                                      X_test, 
                                                                   y_test)
        else:
            classification_text, image_path = capture_output_and_plot(classifier_name,
                                                                      accuracy, 
                                                                      Data_tag, 
                                                                      classification_time,
                                                                      classifier, 
                                                                      X_test, 
                                                                 y_test)
  
            def append_results_to_doc(doc, classification_text, image_path, 
                                      font_name='Times New Roman', font_size=11):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(classification_text)
            
                # Set font properties
                font = run.font
                font.name = font_name
                font.size = Pt(font_size)
            
                # Add image
                max_width_in_inches = 6.0  # You can change this to fit your layout
                doc.add_picture(image_path, width=Inches(max_width_in_inches))
                doc.add_paragraph('\n' + '='*70 + '\n')  # Add separation between results
                
      
            append_results_to_doc(
                doc,
                classification_text,
                image_path,
                font_name='Times New Roman',
                font_size=11
                )
            os.remove('temp_image.png') 
 
doc.save(docx_filename)
    

